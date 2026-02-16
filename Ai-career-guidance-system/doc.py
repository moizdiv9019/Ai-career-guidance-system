from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_roadmap_docx(roadmap_data, filename="Roadmap.docx"):
    doc = Document()

    # Title of the Document
    title = doc.add_heading('Personalized Career Roadmap', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for month_data in roadmap_data:
        # Month Header
        doc.add_heading(f"Month {month_data['month']}: {month_data['focus']}", level=1)
        
        for week in month_data['weeks']:
            # Week Subheader
            doc.add_heading(f"Week {week['week_number']}: {week['week_title']}", level=2)
            
            # Learning Objectives
            doc.add_paragraph("Learning Objectives:", style='List Bullet')
            for obj in week['learning_objectives']:
                p = doc.add_paragraph(obj, style='List Bullet 2')
                p.paragraph_format.left_indent = Pt(40)

            # Tools
            tools_str = ", ".join(week['tools'])
            doc.add_paragraph(f"Tools: {tools_str}", style='No Spacing')

            # Practice Task
            task_para = doc.add_paragraph()
            run = task_para.add_run("Practice Task: ")
            run.bold = True
            task_para.add_run(week['practice_task'])
            
            doc.add_paragraph() # Add spacing

    doc.save(filename)
    return filename


